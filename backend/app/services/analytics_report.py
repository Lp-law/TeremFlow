"""
Client-facing report generation (PDF/DOCX) from Analytics v2 data.
RTL-friendly; no case identifiers. Fixed Light Navy + Gold premium palette.
"""
from __future__ import annotations

import io
import re
from datetime import datetime
from typing import Any

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle, Image

# Optional matplotlib for chart
try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    HAS_MATPLOTLIB = True
except ImportError:
    HAS_MATPLOTLIB = False

# Optional python-docx for DOCX
try:
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


CASE_TYPE_LABELS: dict[str, str] = {
    "COURT": 'תיק ביהמ"ש',
    "DEMAND_LETTER": "מכתב דרישה",
    "SMALL_CLAIMS": "תביעות קטנות",
}

# Premium palette (exact hex)
PRIMARY_NAVY = "#1F3B63"
ACCENT_GOLD = "#C9A227"
HEADER_BG = PRIMARY_NAVY
HEADER_TEXT = "#FFFFFF"
SUBTLE_GRAY = "#F3F5F7"
BODY_TEXT = "#111827"


def _parse_hex(hex_str: str | None) -> tuple[float, float, float]:
    """Parse #RRGGBB to (r,g,b) in 0..1."""
    if not hex_str or not isinstance(hex_str, str):
        hex_str = PRIMARY_NAVY
    hex_str = hex_str.strip()
    m = re.match(r"#?([0-9A-Fa-f]{6})$", hex_str)
    if not m:
        return (31/255, 59/255, 99/255)
    s = m.group(1)
    return (int(s[0:2], 16) / 255.0, int(s[2:4], 16) / 255.0, int(s[4:6], 16) / 255.0)


def _hex_to_reportlab(hex_str: str | None):
    r, g, b = _parse_hex(hex_str)
    return colors.Color(r, g, b)


def _chart_closing_stage_png(distributions: dict, primary_hex: str | None = None, accent_hex: str | None = None) -> bytes | None:
    """Horizontal bar chart: navy bars, max bar gold; axes body_text; minimal grid."""
    if not HAS_MATPLOTLIB:
        return None
    closing = distributions.get("closing_stage") or []
    if not closing:
        return None
    labels = [r.get("label") or r.get("code", "") for r in closing]
    counts = [int(r.get("count", 0)) for r in closing]
    if not labels or not counts:
        return None
    navy = _parse_hex(primary_hex or PRIMARY_NAVY)
    gold = _parse_hex(accent_hex or ACCENT_GOLD)
    body_rgb = _parse_hex(BODY_TEXT)
    max_idx = max(range(len(counts)), key=lambda i: counts[i])
    bar_colors = [gold if i == max_idx else navy for i in range(len(counts))]
    fig, ax = plt.subplots(figsize=(5, max(3, len(labels) * 0.4)))
    y_pos = range(len(labels))
    ax.barh(y_pos, counts, color=bar_colors, height=0.6)
    ax.set_yticks(y_pos)
    ax.set_yticklabels(labels, fontsize=9, color=body_rgb)
    ax.invert_yaxis()
    ax.set_xlabel("תיקים", fontsize=9, color=body_rgb)
    ax.tick_params(axis="x", colors=body_rgb)
    ax.tick_params(axis="y", colors=body_rgb)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.yaxis.set_ticks_position("left")
    ax.xaxis.set_ticks_position("bottom")
    ax.grid(axis="x", alpha=0.15, linestyle="-")
    fig.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format="png", dpi=100, bbox_inches="tight")
    plt.close()
    buf.seek(0)
    return buf.read()


def _narrative_bullets(data: dict) -> list[str]:
    """Generate short factual narrative bullets from v2 data (no case IDs)."""
    bullets: list[str] = []
    filters = data.get("filters") or {}
    kpis = data.get("kpis") or {}
    extra = data.get("extra_metrics") or {}
    totals = data.get("totals") or {}
    denom = int(filters.get("denominator_cases") or 0)
    start = filters.get("start_date") or ""
    end = filters.get("end_date") or ""

    if denom > 0:
        bullets.append(f"במהלך התקופה נכללו {denom} תיקים (תאריך פתיחה {start} עד {end}).")
    avg_stage = kpis.get("avg_stage_fee_ils")
    if avg_stage is not None:
        s = str(avg_stage) if not isinstance(avg_stage, (int, float)) else f"{float(avg_stage):,.2f}"
        bullets.append(f"שכ״ט ממוצע לפי שלבים היה {s} ₪.")
    avg_ret = kpis.get("avg_retainer_fee_ils")
    if avg_ret is not None:
        s = str(avg_ret) if not isinstance(avg_ret, (int, float)) else f"{float(avg_ret):,.2f}"
        bullets.append(f"שכ״ט ממוצע לפי ריטיינר (תיאורטי) היה {s} ₪.")
    idx_denom = int(extra.get("closing_stage_index_denominator_cases") or 0)
    if idx_denom > 0:
        avg_idx = float(extra.get("avg_closing_stage_index") or 0)
        bullets.append(f"השלב השכיח לסיום (תיקי ביהמ״ש סגורים): ממוצע שלב {avg_idx:.2f} (מבוסס על {idx_denom} תיקים בשלבים 1–5).")
    by_branch = totals.get("by_branch") or []
    if by_branch:
        top = max(by_branch, key=lambda r: int(r.get("count") or 0))
        bn = top.get("branch_name") or "ללא סניף"
        cnt = top.get("count") or 0
        bullets.append(f"הסניף עם נפח הגבוה ביותר: {bn} ({cnt} תיקים).")
    return bullets


def _fmt_ils(val: Any) -> str:
    if val is None:
        return "0.00"
    if isinstance(val, (int, float)):
        return f"{float(val):,.2f}"
    return str(val)


def _pdf_table_style_premium(header_bg, header_text, subtle_gray, accent_gold, nrows: int, ncols: int, totals_row: int | None = None):
    """TableStyle: header row navy/white; zebra subtle_gray; light borders; optional totals row bold + gold line."""
    light_gray = colors.HexColor("#E5E7EB")
    styles = [
        ("ALIGN", (0, 0), (-1, -1), "RIGHT"),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("TEXTCOLOR", (0, 0), (-1, 0), header_text),
        ("BACKGROUND", (0, 0), (-1, 0), header_bg),
        ("LINEBELOW", (0, 0), (-1, 0), 1, header_bg),
        ("GRID", (0, 0), (-1, -1), 0.5, light_gray),
    ]
    for r in range(1, nrows):
        if r % 2 == 1:
            styles.append(("BACKGROUND", (0, r), (-1, r), subtle_gray))
    if totals_row is not None and 0 <= totals_row < nrows:
        styles.append(("FONTNAME", (0, totals_row), (-1, totals_row), "Helvetica-Bold"))
        styles.append(("LINEABOVE", (0, totals_row), (-1, totals_row), 1.5, accent_gold))
    return TableStyle(styles)


def _pdf_section_heading(story, text: str, style_heading, accent: Any) -> None:
    """Add section heading (navy) and thin gold underline."""
    story.append(Paragraph(text, style_heading))
    line_tbl = Table([[""]], colWidths=[16 * cm])
    line_tbl.setStyle(TableStyle([("LINEBELOW", (0, 0), (-1, -1), 1.5, accent), ("TOPPADDING", (0, 0), (-1, -1), 2), ("BOTTOMPADDING", (0, 0), (-1, -1), 6)]))
    story.append(line_tbl)


def build_report_pdf(data: dict, template_id: str, brand: dict | None) -> bytes:
    """PDF with full-width title bar (navy/white), section headers (navy + gold underline), zebra tables, chart navy+gold."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=1.5 * cm,
        bottomMargin=1.5 * cm,
    )
    primary_hex = (brand or {}).get("primary_hex") or PRIMARY_NAVY
    accent_hex = (brand or {}).get("accent_hex") or ACCENT_GOLD
    header_bg = _hex_to_reportlab(HEADER_BG)
    header_text = _hex_to_reportlab(HEADER_TEXT)
    primary = _hex_to_reportlab(primary_hex)
    accent = _hex_to_reportlab(accent_hex)
    subtle_gray = _hex_to_reportlab(SUBTLE_GRAY)
    body_color = _hex_to_reportlab(BODY_TEXT)

    styles = getSampleStyleSheet()
    style_rtl = ParagraphStyle(
        name="RTL",
        parent=styles["Normal"],
        alignment=2,
        fontSize=11,
        textColor=body_color,
    )
    style_heading = ParagraphStyle(
        name="HeadingRTL",
        parent=styles["Normal"],
        alignment=2,
        fontSize=12,
        textColor=primary,
        spaceAfter=2,
    )
    story = []

    title = "דו״ח סיכום פעילות"
    if template_id == "T2":
        title = "דו״ח סניפים"
    elif template_id == "T3":
        title = "דו״ח סוגי תיקים"
    filters = data.get("filters") or {}
    period_text = f"תקופה: {filters.get('start_date', '')} – {filters.get('end_date', '')}  |  תיקים בפילטר: {filters.get('denominator_cases', 0)}"
    title_bar = Table([[Paragraph(f"<b>{title}</b><br/><font size=9>{period_text}</font>", ParagraphStyle(name="TitleBar", alignment=2, fontSize=14, textColor=header_text))]], colWidths=[16 * cm])
    title_bar.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), header_bg),
        ("TOPPADDING", (0, 0), (-1, -1), 10),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
    ]))
    story.append(title_bar)
    story.append(Spacer(1, 0.6 * cm))

    kpis = data.get("kpis") or {}
    _pdf_section_heading(story, "<b>מפתחות ביצוע</b>", style_heading, accent)
    story.append(Paragraph(f"שכ״ט ממוצע לפי שלבים: {_fmt_ils(kpis.get('avg_stage_fee_ils'))} ₪", style_rtl))
    story.append(Paragraph(f"שכ״ט ממוצע לפי ריטיינר (תיאורטי): {_fmt_ils(kpis.get('avg_retainer_fee_ils'))} ₪", style_rtl))
    story.append(Paragraph(f"הוצאות ממוצעות לתיק: {_fmt_ils(kpis.get('avg_expenses_ils'))} ₪", style_rtl))
    extra = data.get("extra_metrics") or {}
    if extra.get("closing_stage_index_denominator_cases"):
        story.append(Paragraph(f"שלב סיום ממוצע (תיקי ביהמ״ש): {float(extra.get('avg_closing_stage_index', 0)):.2f} (מבוסס על {extra.get('closing_stage_index_denominator_cases')} תיקים)", style_rtl))
    story.append(Spacer(1, 0.5 * cm))

    dist = data.get("distributions") or {}
    branch_ct = dist.get("branch_case_type") or []
    branch_fee = data.get("branch_fee_averages") or []
    case_type_fee = data.get("case_type_fee_averages") or []
    closing_stage = dist.get("closing_stage") or []
    totals = data.get("totals") or {}
    by_branch = totals.get("by_branch") or []

    if template_id == "T1":
        _pdf_section_heading(story, "<b>סיכום</b>", style_heading, accent)
        for bullet in _narrative_bullets(data):
            story.append(Paragraph(f"• {bullet}", style_rtl))
        story.append(Spacer(1, 0.6 * cm))
        chart_png = _chart_closing_stage_png(dist, primary_hex, accent_hex)
        if chart_png:
            try:
                _pdf_section_heading(story, "<b>התפלגות שלב סיום (תיקים סגורים)</b>", style_heading, accent)
                story.append(Image(io.BytesIO(chart_png), width=12 * cm, height=8 * cm))
                story.append(Spacer(1, 0.5 * cm))
            except Exception:
                pass
        if branch_ct:
            _pdf_section_heading(story, "<b>נפח לפי סניף וסוג תיק</b>", style_heading, accent)
            table_data = [["סניף", "סוג תיק", "כמות"]]
            for r in branch_ct[:15]:
                bn = r.get("branch_name") or "ללא סניף"
                ct = CASE_TYPE_LABELS.get(str(r.get("case_type", "")), str(r.get("case_type", "")))
                table_data.append([bn, ct, str(r.get("count", 0))])
            t = Table(table_data, colWidths=[5 * cm, 5 * cm, 2 * cm])
            t.setStyle(_pdf_table_style_premium(header_bg, header_text, subtle_gray, accent, len(table_data), 3, None))
            story.append(t)
            story.append(Spacer(1, 0.5 * cm))
        if branch_fee:
            _pdf_section_heading(story, "<b>שכ״ט ממוצע לפי סניף</b>", style_heading, accent)
            table_data = [["סניף", "תיקים", "שכ״ט שלבים", "שכ״ט ריטיינר", "הוצאות"]]
            for r in branch_fee[:10]:
                table_data.append([str(r.get("branch_name", "")), str(r.get("cases_count", 0)), _fmt_ils(r.get("avg_stage_fee_ils")), _fmt_ils(r.get("avg_retainer_fee_ils")), _fmt_ils(r.get("avg_expenses_ils"))])
            t = Table(table_data, colWidths=[4 * cm, 2 * cm, 3 * cm, 3 * cm, 3 * cm])
            t.setStyle(_pdf_table_style_premium(header_bg, header_text, subtle_gray, accent, len(table_data), 5, None))
            story.append(t)

    elif template_id == "T2":
        if branch_fee:
            _pdf_section_heading(story, "<b>שכ״ט ממוצע לפי סניף</b>", style_heading, accent)
            table_data = [["סניף", "תיקים", "שכ״ט שלבים", "שכ״ט ריטיינר", "הוצאות"]]
            for r in branch_fee:
                table_data.append([str(r.get("branch_name", "")), str(r.get("cases_count", 0)), _fmt_ils(r.get("avg_stage_fee_ils")), _fmt_ils(r.get("avg_retainer_fee_ils")), _fmt_ils(r.get("avg_expenses_ils"))])
            t = Table(table_data, colWidths=[4 * cm, 2 * cm, 3 * cm, 3 * cm, 3 * cm])
            t.setStyle(_pdf_table_style_premium(header_bg, header_text, subtle_gray, accent, len(table_data), 5, None))
            story.append(t)
            story.append(Spacer(1, 0.5 * cm))
        if branch_ct:
            _pdf_section_heading(story, "<b>נפח תיקים לפי סניף וסוג תיק</b>", style_heading, accent)
            table_data = [["סניף", "סוג תיק", "כמות"]]
            for r in branch_ct[:20]:
                bn = r.get("branch_name") or "ללא סניף"
                ct = CASE_TYPE_LABELS.get(str(r.get("case_type", "")), str(r.get("case_type", "")))
                table_data.append([bn, ct, str(r.get("count", 0))])
            t = Table(table_data, colWidths=[5 * cm, 5 * cm, 2 * cm])
            t.setStyle(_pdf_table_style_premium(header_bg, header_text, subtle_gray, accent, len(table_data), 3, None))
            story.append(t)
            story.append(Spacer(1, 0.5 * cm))
        _pdf_section_heading(story, "<b>תובנות</b>", style_heading, accent)
        if by_branch:
            top_vol = max(by_branch, key=lambda r: int(r.get("count") or 0))
            bn_vol = top_vol.get("branch_name") or "ללא סניף"
            story.append(Paragraph(f"• הסניף עם נפח הגבוה ביותר: {bn_vol} ({top_vol.get('count', 0)} תיקים).", style_rtl))
        if branch_fee:
            top_stage = max(branch_fee, key=lambda r: float(r.get("avg_stage_fee_ils") or 0))
            story.append(Paragraph(f"• הסניף עם שכ״ט ממוצע לפי שלבים הגבוה ביותר: {top_stage.get('branch_name', '')} ({_fmt_ils(top_stage.get('avg_stage_fee_ils'))} ₪).", style_rtl))
            top_ret = max(branch_fee, key=lambda r: float(r.get("avg_retainer_fee_ils") or 0))
            story.append(Paragraph(f"• הסניף עם שכ״ט ריטיינר ממוצע הגבוה ביותר: {top_ret.get('branch_name', '')} ({_fmt_ils(top_ret.get('avg_retainer_fee_ils'))} ₪).", style_rtl))
        story.append(Spacer(1, 0.5 * cm))
        if closing_stage:
            chart_png = _chart_closing_stage_png(dist, primary_hex, accent_hex)
            if chart_png:
                try:
                    _pdf_section_heading(story, "<b>התפלגות שלב סיום (תיקים סגורים)</b>", style_heading, accent)
                    story.append(Image(io.BytesIO(chart_png), width=10 * cm, height=6 * cm))
                except Exception:
                    pass

    else:
        if case_type_fee:
            _pdf_section_heading(story, "<b>סיכום לפי סוג תיק</b>", style_heading, accent)
            table_data = [["סוג תיק", "תיקים", "שכ״ט שלבים", "שכ״ט ריטיינר", "הוצאות"]]
            for r in case_type_fee:
                ct_label = CASE_TYPE_LABELS.get(str(r.get("case_type", "")), str(r.get("case_type", "")))
                table_data.append([ct_label, str(r.get("cases_count", 0)), _fmt_ils(r.get("avg_stage_fee_ils")), _fmt_ils(r.get("avg_retainer_fee_ils")), _fmt_ils(r.get("avg_expenses_ils"))])
            t = Table(table_data, colWidths=[4 * cm, 2 * cm, 3 * cm, 3 * cm, 3 * cm])
            t.setStyle(_pdf_table_style_premium(header_bg, header_text, subtle_gray, accent, len(table_data), 5, None))
            story.append(t)
            story.append(Spacer(1, 0.5 * cm))
        if closing_stage:
            chart_png = _chart_closing_stage_png(dist, primary_hex, accent_hex)
            if chart_png:
                try:
                    _pdf_section_heading(story, "<b>התפלגות שלב סיום (תיקים סגורים)</b>", style_heading, accent)
                    story.append(Image(io.BytesIO(chart_png), width=12 * cm, height=8 * cm))
                    story.append(Spacer(1, 0.5 * cm))
                except Exception:
                    pass
        if by_branch:
            _pdf_section_heading(story, "<b>סיכום סניפים</b>", style_heading, accent)
            table_data = [["סניף", "תיקים"]]
            for r in by_branch[:10]:
                table_data.append([r.get("branch_name") or "ללא סניף", str(r.get("count", 0))])
            t = Table(table_data, colWidths=[8 * cm, 3 * cm])
            t.setStyle(_pdf_table_style_premium(header_bg, header_text, subtle_gray, accent, len(table_data), 2, None))
            story.append(t)

    doc.build(story)
    buf.seek(0)
    return buf.read()


def _docx_set_cell_shading(cell, hex_str: str) -> None:
    """Set table cell background to hex color (python-docx)."""
    if not HAS_DOCX:
        return
    r, g, b = _parse_hex(hex_str)
    # hex for OOXML is RRGGBB
    hex_val = f"{int(r*255):02X}{int(g*255):02X}{int(b*255):02X}"
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_val)
    tcPr.append(shd)


def _docx_header_cell(cell, header_bg_hex: str) -> None:
    """Set cell shading to header bg and text to white."""
    _docx_set_cell_shading(cell, header_bg_hex)
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)


def _docx_heading_paragraph(doc, text: str, primary_hex: str, gold_underline: bool = True):
    """Add heading (navy bold) and optionally thin gold line below."""
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OxmlElement
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    r, g, b = _parse_hex(primary_hex)
    rPr = run._r.get_or_add_rPr()
    color = _OxmlElement("w:color")
    color.set(_qn("w:val"), f"{int(r*255):02X}{int(g*255):02X}{int(b*255):02X}")
    rPr.append(color)
    if gold_underline:
        try:
            pPr = p._p.get_or_add_pPr()
            bottom = _OxmlElement("w:pBdr")
            el = _OxmlElement("w:bottom")
            el.set(_qn("w:val"), "single")
            el.set(_qn("w:sz"), "6")
            el.set(_qn("w:color"), "C9A227")
            bottom.append(el)
            pPr.append(bottom)
        except Exception:
            pass
    return p


def build_report_docx(data: dict, template_id: str, brand: dict | None) -> bytes:
    """DOCX with navy headings, gold underline, navy table headers (white text). No logo. No case identifiers."""
    if not HAS_DOCX:
        raise RuntimeError("python-docx is not installed")
    doc = Document()
    primary_hex = (brand or {}).get("primary_hex") or PRIMARY_NAVY
    header_bg_hex = HEADER_BG

    title = "דו״ח סיכום פעילות"
    if template_id == "T2":
        title = "דו״ח סניפים"
    elif template_id == "T3":
        title = "דו״ח סוגי תיקים"
    p = doc.add_paragraph(title)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(18)
    r, g, b = _parse_hex(primary_hex)
    p.runs[0].font.color.rgb = RGBColor(int(r * 255), int(g * 255), int(b * 255))

    # Optional: thin gold line under title (paragraph border)
    try:
        p_border = p._p.get_or_add_pPr()
        from docx.oxml import OxmlElement as _Oxml
        bottom = _Oxml("w:pBdr")
        el = _Oxml("w:bottom")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:color"), "C9A227")
        bottom.append(el)
        p_border.append(bottom)
    except Exception:
        pass

    filters = data.get("filters") or {}
    p2 = doc.add_paragraph(f"תקופה: {filters.get('start_date', '')} – {filters.get('end_date', '')} | תיקים: {filters.get('denominator_cases', 0)}")
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    doc.add_paragraph()
    _docx_heading_paragraph(doc, "מפתחות ביצוע", primary_hex)
    kpis = data.get("kpis") or {}
    for line in [
        f"שכ״ט ממוצע לפי שלבים: {_fmt_ils(kpis.get('avg_stage_fee_ils'))} ₪",
        f"שכ״ט ממוצע לפי ריטיינר: {_fmt_ils(kpis.get('avg_retainer_fee_ils'))} ₪",
        f"הוצאות ממוצעות לתיק: {_fmt_ils(kpis.get('avg_expenses_ils'))} ₪",
    ]:
        pr = doc.add_paragraph(line)
        pr.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    dist = data.get("distributions") or {}
    branch_fee = data.get("branch_fee_averages") or []
    branch_ct = dist.get("branch_case_type") or []
    case_type_fee = data.get("case_type_fee_averages") or []
    closing_stage = dist.get("closing_stage") or []

    if template_id == "T1":
        doc.add_paragraph()
        _docx_heading_paragraph(doc, "סיכום", primary_hex)
        for bullet in _narrative_bullets(data):
            pr = doc.add_paragraph(bullet)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        chart_png = _chart_closing_stage_png(dist, primary_hex, ACCENT_GOLD)
        if chart_png:
            doc.add_paragraph()
            _docx_heading_paragraph(doc, "התפלגות שלב סיום", primary_hex)
            doc.add_picture(io.BytesIO(chart_png), width=Cm(12))
        if branch_ct:
            doc.add_paragraph()
            _docx_heading_paragraph(doc, "נפח לפי סניף וסוג תיק", primary_hex)
            table = doc.add_table(rows=1 + len(branch_ct[:15]), cols=3)
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = "סניף", "סוג תיק", "כמות"
            for j in range(3):
                _docx_header_cell(table.rows[0].cells[j], header_bg_hex)
            for i, r in enumerate(branch_ct[:15]):
                row = table.rows[i + 1]
                row.cells[0].text = r.get("branch_name") or "ללא סניף"
                row.cells[1].text = CASE_TYPE_LABELS.get(str(r.get("case_type", "")), str(r.get("case_type", "")))
                row.cells[2].text = str(r.get("count", 0))
        if branch_fee:
            doc.add_paragraph()
            _docx_heading_paragraph(doc, "שכ״ט ממוצע לפי סניף", primary_hex)
            table = doc.add_table(rows=1 + len(branch_fee[:10]), cols=5)
            table.rows[0].cells[0].text = "סניף"
            table.rows[0].cells[1].text = "תיקים"
            table.rows[0].cells[2].text = "שכ״ט שלבים"
            table.rows[0].cells[3].text = "שכ״ט ריטיינר"
            table.rows[0].cells[4].text = "הוצאות"
            for j in range(5):
                _docx_header_cell(table.rows[0].cells[j], header_bg_hex)
            for i, r in enumerate(branch_fee[:10]):
                row = table.rows[i + 1]
                row.cells[0].text = str(r.get("branch_name", ""))
                row.cells[1].text = str(r.get("cases_count", 0))
                row.cells[2].text = _fmt_ils(r.get("avg_stage_fee_ils"))
                row.cells[3].text = _fmt_ils(r.get("avg_retainer_fee_ils"))
                row.cells[4].text = _fmt_ils(r.get("avg_expenses_ils"))

    elif template_id == "T2":
        if branch_fee:
            doc.add_paragraph()
            _docx_heading_paragraph(doc, "שכ״ט ממוצע לפי סניף", primary_hex)
            table = doc.add_table(rows=1 + len(branch_fee), cols=5)
            table.rows[0].cells[0].text = "סניף"
            table.rows[0].cells[1].text = "תיקים"
            table.rows[0].cells[2].text = "שכ״ט שלבים"
            table.rows[0].cells[3].text = "שכ״ט ריטיינר"
            table.rows[0].cells[4].text = "הוצאות"
            for j in range(5):
                _docx_header_cell(table.rows[0].cells[j], header_bg_hex)
            for i, r in enumerate(branch_fee):
                row = table.rows[i + 1]
                row.cells[0].text = str(r.get("branch_name", ""))
                row.cells[1].text = str(r.get("cases_count", 0))
                row.cells[2].text = _fmt_ils(r.get("avg_stage_fee_ils"))
                row.cells[3].text = _fmt_ils(r.get("avg_retainer_fee_ils"))
                row.cells[4].text = _fmt_ils(r.get("avg_expenses_ils"))
        if branch_ct:
            doc.add_paragraph()
            _docx_heading_paragraph(doc, "נפח תיקים לפי סניף וסוג תיק", primary_hex)
            table = doc.add_table(rows=1 + min(20, len(branch_ct)), cols=3)
            table.rows[0].cells[0].text, table.rows[0].cells[1].text, table.rows[0].cells[2].text = "סניף", "סוג תיק", "כמות"
            for j in range(3):
                _docx_header_cell(table.rows[0].cells[j], header_bg_hex)
            for i, r in enumerate(branch_ct[:20]):
                row = table.rows[i + 1]
                row.cells[0].text = r.get("branch_name") or "ללא סניף"
                row.cells[1].text = CASE_TYPE_LABELS.get(str(r.get("case_type", "")), str(r.get("case_type", "")))
                row.cells[2].text = str(r.get("count", 0))
        totals = data.get("totals") or {}
        by_branch = totals.get("by_branch") or []
        if by_branch or branch_fee:
            doc.add_paragraph()
            _docx_heading_paragraph(doc, "תובנות", primary_hex)
            if by_branch:
                top_vol = max(by_branch, key=lambda r: int(r.get("count") or 0))
                pr = doc.add_paragraph(f"• סניף נפח גבוה: {top_vol.get('branch_name') or 'ללא סניף'} ({top_vol.get('count', 0)} תיקים).")
                pr.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            if branch_fee:
                top_stage = max(branch_fee, key=lambda r: float(r.get("avg_stage_fee_ils") or 0))
                pr = doc.add_paragraph(f"• סניף שכ״ט שלבים ממוצע גבוה: {top_stage.get('branch_name', '')} ({_fmt_ils(top_stage.get('avg_stage_fee_ils'))} ₪).")
                pr.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                top_ret = max(branch_fee, key=lambda r: float(r.get("avg_retainer_fee_ils") or 0))
                pr = doc.add_paragraph(f"• סניף שכ״ט ריטיינר ממוצע גבוה: {top_ret.get('branch_name', '')} ({_fmt_ils(top_ret.get('avg_retainer_fee_ils'))} ₪).")
                pr.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        if closing_stage:
            chart_png = _chart_closing_stage_png(dist, primary_hex, ACCENT_GOLD)
            if chart_png:
                doc.add_paragraph()
                _docx_heading_paragraph(doc, "התפלגות שלב סיום", primary_hex)
                doc.add_picture(io.BytesIO(chart_png), width=Cm(10))

    else:
        # T3
        if case_type_fee:
            doc.add_paragraph()
            _docx_heading_paragraph(doc, "סיכום לפי סוג תיק", primary_hex)
            table = doc.add_table(rows=1 + len(case_type_fee), cols=5)
            table.rows[0].cells[0].text = "סוג תיק"
            table.rows[0].cells[1].text = "תיקים"
            table.rows[0].cells[2].text = "שכ״ט שלבים"
            table.rows[0].cells[3].text = "שכ״ט ריטיינר"
            table.rows[0].cells[4].text = "הוצאות"
            for j in range(5):
                _docx_header_cell(table.rows[0].cells[j], header_bg_hex)
            for i, r in enumerate(case_type_fee):
                row = table.rows[i + 1]
                row.cells[0].text = CASE_TYPE_LABELS.get(str(r.get("case_type", "")), str(r.get("case_type", "")))
                row.cells[1].text = str(r.get("cases_count", 0))
                row.cells[2].text = _fmt_ils(r.get("avg_stage_fee_ils"))
                row.cells[3].text = _fmt_ils(r.get("avg_retainer_fee_ils"))
                row.cells[4].text = _fmt_ils(r.get("avg_expenses_ils"))
        _docx_heading_paragraph(doc, "סיכום", primary_hex)
        for bullet in _narrative_bullets(data):
            pr = doc.add_paragraph(bullet)
            pr.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        if closing_stage:
            chart_png = _chart_closing_stage_png(dist, primary_hex, ACCENT_GOLD)
            if chart_png:
                doc.add_paragraph()
                _docx_heading_paragraph(doc, "התפלגות שלב סיום", primary_hex)
                doc.add_picture(io.BytesIO(chart_png), width=Cm(12))
        by_branch = (data.get("totals") or {}).get("by_branch") or []
        if by_branch:
            doc.add_paragraph()
            _docx_heading_paragraph(doc, "סיכום סניפים", primary_hex)
            table = doc.add_table(rows=1 + min(10, len(by_branch)), cols=2)
            table.rows[0].cells[0].text = "סניף"
            table.rows[0].cells[1].text = "תיקים"
            for j in range(2):
                _docx_header_cell(table.rows[0].cells[j], header_bg_hex)
            for i, r in enumerate(by_branch[:10]):
                row = table.rows[i + 1]
                row.cells[0].text = r.get("branch_name") or "ללא סניף"
                row.cells[1].text = str(r.get("count", 0))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def build_client_report(
    data: dict,
    template_id: str,
    report_format: str,
    brand: dict | None,
) -> tuple[bytes, str]:
    """
    Build report bytes and suggested filename.
    data: Analytics v2 response as dict (from model_dump(mode="json")).
    Returns (content_bytes, filename).
    """
    date_str = datetime.now().strftime("%Y%m%d")
    if report_format.lower() == "docx":
        content = build_report_docx(data, template_id, brand)
        ext = "docx"
    else:
        content = build_report_pdf(data, template_id, brand)
        ext = "pdf"
    filename = f"client_report_{template_id}_{date_str}.{ext}"
    return (content, filename)
