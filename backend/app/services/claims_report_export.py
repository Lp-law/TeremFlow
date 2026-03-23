from __future__ import annotations

import io
import logging
from collections import defaultdict
from datetime import date, datetime
from decimal import Decimal
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

try:
    from docxtpl import DocxTemplate
except Exception:  # pragma: no cover
    DocxTemplate = None  # type: ignore[assignment]

from app.models.enums import ClaimsCategory, ClaimsFinalOutcomeType, ClaimsReportCaseStatus
from app.services.claims_reports import compute_default_narrative

logger = logging.getLogger(__name__)

CATEGORY_ORDER = [
    ClaimsCategory.COURT_REPORTED_TO_INSURER,
    ClaimsCategory.REPORTED_WITHOUT_CLAIM,
    ClaimsCategory.NOT_REPORTED_TO_INSURER,
    ClaimsCategory.NON_MEDICAL_MALPRACTICE,
    ClaimsCategory.OTHER,
]

CATEGORY_LABELS: dict[ClaimsCategory, str] = {
    ClaimsCategory.COURT_REPORTED_TO_INSURER: "תיקים המתנהלים בבתי משפט ודווחו לחברת הביטוח",
    ClaimsCategory.REPORTED_WITHOUT_CLAIM: "מקרים שדווחו לחברת הביטוח ללא תביעה",
    ClaimsCategory.NOT_REPORTED_TO_INSURER: "מקרים שלא דווחו לחברת הביטוח",
    ClaimsCategory.NON_MEDICAL_MALPRACTICE: "תיקים שאינם בתחום הרשלנות הרפואית",
    ClaimsCategory.OTHER: "אחר",
}

STATUS_LABELS: dict[ClaimsReportCaseStatus, str] = {
    ClaimsReportCaseStatus.OPEN: "פתוח",
    ClaimsReportCaseStatus.CLOSED: "סגור",
    ClaimsReportCaseStatus.CANNOT_ASSESS_YET: "לא ניתן להעריך בשלב זה",
    ClaimsReportCaseStatus.NO_EXPOSURE: "ללא חשיפה",
    ClaimsReportCaseStatus.REJECTED_EXPECTED: "צפי לדחייה",
    ClaimsReportCaseStatus.SETTLED: "פשרה",
    ClaimsReportCaseStatus.JUDGMENT: "פסק דין",
    ClaimsReportCaseStatus.REJECTED: "נדחה",
    ClaimsReportCaseStatus.REJECTED_WITH_COSTS: "נדחה עם הוצאות",
}

OUTCOME_LABELS: dict[ClaimsFinalOutcomeType, str] = {
    ClaimsFinalOutcomeType.SETTLEMENT: "פשרה",
    ClaimsFinalOutcomeType.JUDGMENT_FOR_PLAINTIFF: "פסק דין לטובת התובע",
    ClaimsFinalOutcomeType.CLAIM_REJECTED: "תביעה נדחתה",
    ClaimsFinalOutcomeType.CLAIM_REJECTED_WITH_COSTS: "תביעה נדחתה עם הוצאות",
    ClaimsFinalOutcomeType.CLOSED_WITHOUT_PAYMENT: "נסגר ללא תשלום",
    ClaimsFinalOutcomeType.OTHER: "אחר",
}

TEMPLATE_ROOT = Path(__file__).resolve().parent.parent / "templates" / "claims_reports"
DEFAULT_TEMPLATE_FILE = TEMPLATE_ROOT / "terem_claims_report_template.docx"
HEBREW_MONTHS = {
    1: "ינואר",
    2: "פברואר",
    3: "מרץ",
    4: "אפריל",
    5: "מאי",
    6: "יוני",
    7: "יולי",
    8: "אוגוסט",
    9: "ספטמבר",
    10: "אוקטובר",
    11: "נובמבר",
    12: "דצמבר",
}
NUMBERED_CATEGORY_TITLES = {
    1: "1. תיקים שעניינם מתנהלים כיום בבתי המשפט השונים ודווחו לחברת הביטוח:",
    2: "2. מקרים שדווחו לחברת הביטוח אך לא הוגשה בהם תביעה:",
    3: "3. מקרים שלא דווחו לחברת הביטוח:",
    4: "4. תיקים שאינם בתחום הרשלנות הרפואית:",
}


def _fmt_ils(v: Any) -> str:
    if v is None:
        return "—"
    try:
        d = Decimal(str(v)).quantize(Decimal("0.01"))
        return f"{d} ₪"
    except Exception:
        return str(v)


def _fmt_number_plain(v: Any) -> str:
    if v is None:
        return "0"
    try:
        d = Decimal(str(v)).quantize(Decimal("0.01"))
        if d == d.to_integral():
            return f"{int(d):,}"
        return f"{d:,.2f}"
    except Exception:
        return str(v)


def _fmt_date_numeric(v: date | None) -> str:
    if not v:
        return ""
    return v.strftime("%d.%m.%y")


def _fmt_date_hebrew(v: date | None, *, with_comma: bool) -> str:
    if not v:
        return ""
    month_name = HEBREW_MONTHS.get(v.month, "")
    if with_comma:
        return f"{v.day} ב{month_name}, {v.year}"
    return f"{v.day} ב{month_name} {v.year}"


def _build_rendered_line(row, *, category_number: int, row_number: int) -> str:
    """
    Build a single legal-style narrative sentence for template `row.rendered_line`.
    Includes hierarchical numbering (e.g. 1.1, 2.3).
    """
    prefix = f"{category_number}.{row_number}"
    case_name = getattr(row, "case_title", None) or getattr(row, "case_reference_text", None) or f"רשומה {getattr(row, 'id', '')}"
    report_case_status = getattr(row, "report_case_status", None)
    status_label = STATUS_LABELS.get(report_case_status, "לא סווג")
    def _clean(s: str) -> str:
        return s.strip().rstrip(". ")

    parts: list[str] = [f"{prefix} {case_name}"]
    parts.append(f"בסטטוס {status_label}")
    final_outcome_type = getattr(row, "final_outcome_type", None)
    if final_outcome_type is not None:
        parts.append(f"תוצאת ההליך: {OUTCOME_LABELS.get(final_outcome_type, str(final_outcome_type))}")
    final_outcome_amount_ils = getattr(row, "final_outcome_amount_ils", None)
    if final_outcome_amount_ils is not None:
        parts.append(f"סכום סיום ההליך עומד על {_fmt_ils(final_outcome_amount_ils)}")
    current_risk_assessment_ils = getattr(row, "current_risk_assessment_ils", None)
    exposure_for_reserve_ils = getattr(row, "exposure_for_reserve_ils", None)
    if current_risk_assessment_ils is not None:
        parts.append(f"הערכת הסיכון העדכנית היא {_fmt_ils(current_risk_assessment_ils)}")
    elif exposure_for_reserve_ils is not None:
        parts.append(f"החשיפה לצורך הפרשה מוערכת בכ-{_fmt_ils(exposure_for_reserve_ils)}")
    remaining_deductible_ils = getattr(row, "remaining_deductible_ils", None)
    if remaining_deductible_ils is not None:
        parts.append(f"יתרת ההשתתפות העצמית נאמדת בכ-{_fmt_ils(remaining_deductible_ils)}")
    narrative_text = getattr(row, "narrative_text", None)
    if narrative_text:
        parts.append(_clean(str(narrative_text)))
    else:
        parts.append(_clean(compute_default_narrative(row)))
    legal_summary_text = getattr(row, "legal_summary_text", None)
    if legal_summary_text:
        parts.append(f"בהיבט המשפטי: {_clean(str(legal_summary_text))}")
    return ". ".join([_clean(p) for p in parts if p]).strip() + "."


def _set_paragraph_rtl(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    p_pr.append(bidi)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    _set_paragraph_rtl(p)


def _add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    _set_paragraph_rtl(p)
    run = p.add_run(text)
    run.bold = bold


def _add_template_paragraph(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    _set_paragraph_rtl(p)
    run = p.add_run(text)
    run.bold = bold


def ensure_default_template_exists() -> Path:
    TEMPLATE_ROOT.mkdir(parents=True, exist_ok=True)
    if DEFAULT_TEMPLATE_FILE.exists():
        return DEFAULT_TEMPLATE_FILE

    doc = Document()
    _add_template_paragraph(doc, "{{ report_title }}", bold=True)
    _add_template_paragraph(doc, "לקוח/מוסד: {{ client_name }}")
    _add_template_paragraph(doc, "תאריך חתך: {{ report_cutoff_date }}")
    _add_template_paragraph(doc, "מעודכן ליום: {{ updated_to_date }}")
    _add_template_paragraph(doc, "סטטוס דו\"ח: {{ report_status_label }}")
    _add_template_paragraph(doc, "המלצת הפרשה: {{ recommended_reserve_ils }}", bold=True)
    _add_template_paragraph(doc, "")
    _add_template_paragraph(doc, "{{ intro_text }}")
    _add_template_paragraph(doc, "")
    _add_template_paragraph(doc, "{% for category in categories %}", bold=True)
    _add_template_paragraph(doc, "{{ category.title }}", bold=True)
    _add_template_paragraph(doc, "{% for row in category.rows %}")
    _add_template_paragraph(doc, "{{ row.idx }}. {{ row.case_title }}", bold=True)
    _add_template_paragraph(doc, "מספר הליך/תיק: {{ row.proceeding_or_reference }}")
    _add_template_paragraph(doc, "סטטוס בדו\"ח: {{ row.report_status_label }}")
    _add_template_paragraph(doc, "{{ row.final_outcome_line }}")
    _add_template_paragraph(doc, "{{ row.final_amount_line }}")
    _add_template_paragraph(doc, "{{ row.awarded_costs_line }}")
    _add_template_paragraph(doc, "השתתפות עצמית מלאה: {{ row.deductible_ils_gross }}")
    _add_template_paragraph(doc, "שולם על חשבון: {{ row.amount_paid_ils }}")
    _add_template_paragraph(doc, "נותר: {{ row.remaining_deductible_ils }}")
    _add_template_paragraph(doc, "הוצאות: {{ row.expenses_total_ils }}")
    _add_template_paragraph(doc, "שכ\"ט: {{ row.fees_total_ils }}")
    _add_template_paragraph(doc, "חשיפה לרזרבה: {{ row.exposure_for_reserve_ils }}")
    _add_template_paragraph(doc, "נרטיב: {{ row.narrative }}")
    _add_template_paragraph(doc, "{{ row.legal_summary_line }}")
    _add_template_paragraph(doc, "{{ row.status_note_line }}")
    _add_template_paragraph(doc, "{{ row.internal_notes_line }}")
    _add_template_paragraph(doc, "{% endfor %}")
    _add_template_paragraph(doc, "{% endfor %}")
    _add_template_paragraph(doc, "")
    _add_template_paragraph(doc, "{{ closing_text }}")
    doc.save(DEFAULT_TEMPLATE_FILE)
    return DEFAULT_TEMPLATE_FILE


def _resolve_template_path(template_key: str | None) -> Path:
    ensure_default_template_exists()
    if not template_key:
        return DEFAULT_TEMPLATE_FILE
    key = template_key.strip()
    if not key:
        return DEFAULT_TEMPLATE_FILE
    candidate = TEMPLATE_ROOT / f"{key}.docx"
    if candidate.exists():
        return candidate
    return DEFAULT_TEMPLATE_FILE


def _row_context(row, idx: int) -> dict[str, str]:
    outcome_label = OUTCOME_LABELS.get(row.final_outcome_type) if row.final_outcome_type else None
    return {
        "idx": str(idx),
        "case_title": row.case_title or row.case_reference_text or f"רשומה {row.id}",
        "proceeding_or_reference": row.proceeding_number or row.case_reference_text or "—",
        "report_status_label": STATUS_LABELS.get(row.report_case_status, str(row.report_case_status)),
        "final_outcome_line": f"תוצאת סיום: {outcome_label}" if outcome_label else "",
        "final_amount_line": f"סכום סיום: {_fmt_ils(row.final_outcome_amount_ils)}" if row.final_outcome_amount_ils is not None else "",
        "awarded_costs_line": (
            f"הוצאות שנפסקו לטובת טרם: {_fmt_ils(row.awarded_costs_to_terem_ils)}"
            if row.awarded_costs_to_terem_ils is not None
            else ""
        ),
        "deductible_ils_gross": _fmt_ils(row.deductible_ils_gross),
        "amount_paid_ils": _fmt_ils(row.amount_already_paid_on_deductible_ils),
        "remaining_deductible_ils": _fmt_ils(row.remaining_deductible_ils),
        "expenses_total_ils": _fmt_ils(row.expenses_total_ils),
        "fees_total_ils": _fmt_ils(row.fees_total_ils),
        "exposure_for_reserve_ils": _fmt_ils(row.exposure_for_reserve_ils),
        "narrative": row.narrative_text or compute_default_narrative(row),
        "legal_summary_line": f"סיכום משפטי: {row.legal_summary_text}" if row.legal_summary_text else "",
        "status_note_line": f"הערת סטטוס: {row.status_note}" if row.status_note else "",
        "internal_notes_line": f"הערות פנימיות: {row.internal_notes}" if row.internal_notes else "",
    }


def _build_template_context(report, rows: list) -> dict[str, Any]:
    grouped = defaultdict(list)
    for row in rows:
        grouped[row.category_for_report].append(row)

    # Keep strict mapping for 4 template categories.
    category_mapping: list[tuple[int, ClaimsCategory]] = [
        (1, ClaimsCategory.COURT_REPORTED_TO_INSURER),
        (2, ClaimsCategory.REPORTED_WITHOUT_CLAIM),
        (3, ClaimsCategory.NOT_REPORTED_TO_INSURER),
        (4, ClaimsCategory.NON_MEDICAL_MALPRACTICE),
    ]

    category_rows_map: dict[int, list] = {}
    for idx, cat in category_mapping:
        category_rows_map[idx] = list(grouped.get(cat, []))
    # Keep OTHER additive by appending into category 4 so no row is lost.
    if grouped.get(ClaimsCategory.OTHER):
        category_rows_map[4].extend(grouped.get(ClaimsCategory.OTHER, []))

    # Template-specific rows: each row has only `rendered_line` as requested.
    category_template_rows: dict[int, list[dict[str, str]]] = {}
    for idx, _ in category_mapping:
        rows_for_cat = category_rows_map[idx]
        category_template_rows[idx] = [
            {"rendered_line": _build_rendered_line(r, category_number=idx, row_number=row_idx + 1)}
            for row_idx, r in enumerate(rows_for_cat)
        ]

    categories: list[dict[str, Any]] = []
    for idx, cat in category_mapping:
        cat_rows = category_rows_map[idx]
        if not cat_rows:
            continue
        categories.append(
            {
                "key": str(cat),
                "title": NUMBERED_CATEGORY_TITLES[idx],
                "rows": [_row_context(r, idx + 1) for idx, r in enumerate(cat_rows)],
            }
        )

    cutoff = report.report_cutoff_date if getattr(report, "report_cutoff_date", None) else None
    updated_to = report.updated_to_date if getattr(report, "updated_to_date", None) else None
    today = datetime.now().date()
    return {
        # New placeholders requested by active template
        "report_date_hebrew": _fmt_date_hebrew(today, with_comma=False),
        "report_cutoff_date": _fmt_date_numeric(cutoff),
        "updated_to_date": _fmt_date_numeric(updated_to),
        "report_cutoff_date_hebrew": _fmt_date_hebrew(cutoff, with_comma=True),
        "report_cutoff_date_hebrew_compact": _fmt_date_hebrew(cutoff, with_comma=False),
        "reserve_recommendation_amount": _fmt_number_plain(report.recommended_reserve_ils),
        "category_1_title": NUMBERED_CATEGORY_TITLES[1],
        "category_2_title": NUMBERED_CATEGORY_TITLES[2],
        "category_3_title": NUMBERED_CATEGORY_TITLES[3],
        "category_4_title": NUMBERED_CATEGORY_TITLES[4],
        "category_1_rows": category_template_rows[1],
        "category_2_rows": category_template_rows[2],
        "category_3_rows": category_template_rows[3],
        "category_4_rows": category_template_rows[4],
        # Legacy/additional placeholders preserved for compatibility
        "report_id": report.id,
        "report_title": report.title or "דו\"ח תביעות / חשיפות",
        "client_name": report.client_name or "",
        "report_cutoff_date_iso": cutoff.isoformat() if cutoff else "",
        "updated_to_date_iso": updated_to.isoformat() if updated_to else "",
        "report_status_label": "סופי" if str(report.status) == "FINAL" else "טיוטה",
        "recommended_reserve_ils": _fmt_ils(report.recommended_reserve_ils),
        "intro_text": report.intro_text or "",
        "closing_text": report.closing_text or "",
        "categories": categories,
    }


def _build_claims_report_docx_direct(report, rows: list) -> tuple[bytes, str]:
    doc = Document()
    _add_heading(doc, report.title or "דו\"ח תביעות / חשיפות", 0)
    _add_paragraph(doc, f"לקוח/מוסד: {report.client_name}")
    _add_paragraph(doc, f"תאריך חתך: {report.report_cutoff_date.isoformat()}")
    if report.updated_to_date is not None:
        _add_paragraph(doc, f"מעודכן ליום: {report.updated_to_date.isoformat()}")
    _add_paragraph(doc, f"סטטוס דו\"ח: {'סופי' if str(report.status) == 'FINAL' else 'טיוטה'}")
    if report.recommended_reserve_ils is not None:
        _add_paragraph(doc, f"המלצת הפרשה: {_fmt_ils(report.recommended_reserve_ils)}", bold=True)
    if report.intro_text:
        _add_paragraph(doc, "")
        _add_paragraph(doc, report.intro_text)

    grouped = defaultdict(list)
    for row in rows:
        grouped[row.category_for_report].append(row)

    for category in CATEGORY_ORDER:
        category_rows = grouped.get(category, [])
        if not category_rows:
            continue
        _add_paragraph(doc, "")
        _add_heading(doc, CATEGORY_LABELS.get(category, str(category)), 1)
        for idx, row in enumerate(category_rows, start=1):
            row_title = row.case_title or row.case_reference_text or f"רשומה {row.id}"
            _add_paragraph(doc, f"{idx}. {row_title}", bold=True)
            _add_paragraph(doc, f"מספר הליך/תיק: {row.proceeding_number or row.case_reference_text or '—'}")
            _add_paragraph(doc, f"סטטוס בדו\"ח: {STATUS_LABELS.get(row.report_case_status, str(row.report_case_status))}")
            if row.final_outcome_type is not None:
                _add_paragraph(doc, f"תוצאת סיום: {OUTCOME_LABELS.get(row.final_outcome_type, str(row.final_outcome_type))}")
            if row.final_outcome_amount_ils is not None:
                _add_paragraph(doc, f"סכום סיום: {_fmt_ils(row.final_outcome_amount_ils)}")
            if row.awarded_costs_to_terem_ils is not None:
                _add_paragraph(doc, f"הוצאות שנפסקו לטובת טרם: {_fmt_ils(row.awarded_costs_to_terem_ils)}")
            _add_paragraph(doc, f"השתתפות עצמית מלאה: {_fmt_ils(row.deductible_ils_gross)}")
            _add_paragraph(doc, f"שולם על חשבון: {_fmt_ils(row.amount_already_paid_on_deductible_ils)}")
            _add_paragraph(doc, f"נותר: {_fmt_ils(row.remaining_deductible_ils)}")
            _add_paragraph(doc, f"הוצאות: {_fmt_ils(row.expenses_total_ils)}")
            _add_paragraph(doc, f"שכ\"ט: {_fmt_ils(row.fees_total_ils)}")
            _add_paragraph(doc, f"חשיפה לרזרבה: {_fmt_ils(row.exposure_for_reserve_ils)}")
            narrative = row.narrative_text or compute_default_narrative(row)
            _add_paragraph(doc, f"נרטיב: {narrative}")
            if row.legal_summary_text:
                _add_paragraph(doc, f"סיכום משפטי: {row.legal_summary_text}")
            if row.status_note:
                _add_paragraph(doc, f"הערת סטטוס: {row.status_note}")
            if row.internal_notes:
                _add_paragraph(doc, f"הערות פנימיות: {row.internal_notes}")

    if report.closing_text:
        _add_paragraph(doc, "")
        _add_paragraph(doc, report.closing_text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    file_date = datetime.now().strftime("%Y%m%d")
    filename = f"claims_report_{report.id}_{file_date}.docx"
    return buf.read(), filename


def build_claims_report_docx_from_template(report, rows: list, *, template_key: str | None = None) -> tuple[bytes, str]:
    if DocxTemplate is None:
        raise RuntimeError("docxtpl dependency is not installed")
    template_path = _resolve_template_path(template_key or getattr(report, "template_key", None))
    tpl = DocxTemplate(str(template_path))
    tpl.render(_build_template_context(report, rows))
    buf = io.BytesIO()
    tpl.save(buf)
    buf.seek(0)
    file_date = datetime.now().strftime("%Y%m%d")
    filename = f"claims_report_{report.id}_{file_date}.docx"
    return buf.read(), filename


def build_claims_report_docx(report, rows: list) -> tuple[bytes, str]:
    try:
        return build_claims_report_docx_from_template(report, rows, template_key=getattr(report, "template_key", None))
    except Exception:
        logger.exception("Template export failed; using legacy direct builder")
        return _build_claims_report_docx_direct(report, rows)
